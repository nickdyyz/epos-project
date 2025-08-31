#!/usr/bin/env python3
"""
extract_word_text.py - Extract text from Word (.docx) files

This script processes Word documents from an input directory and extracts
their text content. The extracted text is saved to an output directory
with the same filename but with a .txt extension.

Usage:
    python extract_word_text.py --input <input_dir> --output <output_dir>

Requirements:
    - python-docx library
    - Python 3.6+
"""

import argparse
import logging
import os
import sys
from pathlib import Path
from typing import List, Optional, Tuple

try:
    from docx import Document
except ImportError:
    print("Error: python-docx library not installed.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)


def setup_logging(log_level: str = "INFO") -> None:
    """Set up logging configuration.
    
    Args:
        log_level: The logging level (default: INFO)
    """
    numeric_level = getattr(logging, log_level.upper(), None)
    if not isinstance(numeric_level, int):
        raise ValueError(f"Invalid log level: {log_level}")
    
    logging.basicConfig(
        level=numeric_level,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler(os.path.join(os.path.dirname(__file__), "word_extraction.log"))
        ]
    )


def get_word_files(input_dir: str) -> List[Path]:
    """Get all Word files from the input directory.
    
    Args:
        input_dir: Path to the directory containing Word files
    
    Returns:
        List of Path objects for each Word file
    """
    input_path = Path(input_dir)
    if not input_path.is_dir():
        raise ValueError(f"Input directory does not exist: {input_dir}")
    
    word_files = list(input_path.glob("*.docx"))
    logging.info(f"Found {len(word_files)} Word files in {input_dir}")
    return word_files


def extract_text_from_word(word_path: Path, output_dir: Path) -> Tuple[bool, Optional[str]]:
    """Extract text from a Word file.
    
    Args:
        word_path: Path to the Word file
        output_dir: Directory where the text file will be saved
    
    Returns:
        Tuple containing success status (bool) and error message (str) if any
    """
    output_file = output_dir / f"{word_path.stem}.txt"
    
    try:
        # Ensure output directory exists
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Load the Word document
        logging.debug(f"Processing {word_path}")
        doc = Document(str(word_path))
        
        # Extract text from all paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        # Join all text with newlines
        extracted_text = '\n'.join(full_text)
        
        # Write to output file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(extracted_text)
        
        if output_file.exists():
            logging.info(f"Successfully extracted text from {word_path.name} to {output_file}")
            return True, None
        else:
            return False, "Output file was not created"
            
    except Exception as e:
        error_msg = f"Error processing {word_path.name}: {str(e)}"
        logging.error(error_msg)
        return False, error_msg


def process_word_files(input_dir: str, output_dir: str) -> Tuple[int, int]:
    """Process all Word files in the input directory.
    
    Args:
        input_dir: Directory containing Word files
        output_dir: Directory where text files will be saved
    
    Returns:
        Tuple containing count of successful and failed extractions
    """
    word_files = get_word_files(input_dir)
    output_path = Path(output_dir)
    
    successful = 0
    failed = 0
    
    for word_file in word_files:
        logging.info(f"Processing {word_file.name}")
        success, error = extract_text_from_word(word_file, output_path)
        
        if success:
            successful += 1
        else:
            failed += 1
            logging.error(f"Failed to extract text from {word_file.name}: {error}")
    
    return successful, failed


def parse_arguments() -> argparse.Namespace:
    """Parse command line arguments.
    
    Returns:
        Parsed arguments namespace
    """
    parser = argparse.ArgumentParser(
        description="Extract text from Word (.docx) files"
    )
    parser.add_argument(
        "--input",
        "-i",
        required=True,
        help="Directory containing Word files"
    )
    parser.add_argument(
        "--output",
        "-o", 
        required=True,
        help="Directory where text files will be saved"
    )
    parser.add_argument(
        "--log-level",
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        default="INFO",
        help="Set the logging level"
    )
    
    return parser.parse_args()


def main() -> int:
    """Main function to orchestrate the Word text extraction process.
    
    Returns:
        Exit code (0 for success, 1 for failure)
    """
    args = parse_arguments()
    setup_logging(args.log_level)
    
    logging.info(f"Starting Word text extraction")
    logging.info(f"Input directory: {args.input}")
    logging.info(f"Output directory: {args.output}")
    
    try:
        successful, failed = process_word_files(args.input, args.output)
        
        logging.info(f"Word extraction complete")
        logging.info(f"Successfully processed: {successful}")
        logging.info(f"Failed to process: {failed}")
        
        if failed > 0:
            logging.warning(f"Some Word files could not be processed. Check the log for details.")
            return 1
        return 0
        
    except Exception as e:
        logging.critical(f"Fatal error: {str(e)}")
        return 1


if __name__ == "__main__":
    sys.exit(main())
